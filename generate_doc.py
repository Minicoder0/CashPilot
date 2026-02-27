import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_executive_summary():
    doc = Document()
    
    # Title
    title = doc.add_heading('CashPilot — Executive Summary & Business Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1
    doc.add_heading('1. The Problem & Our Solution', level=1)
    
    p = doc.add_paragraph()
    p.add_run('The Problem: ').bold = True
    p.add_run('Micro-business owners (freelancers, creators, local shops) are overwhelmed by traditional accounting software. Platforms like QuickBooks are built for accountants, requiring knowledge of "charts of accounts," "reconciliation," and "double-entry bookkeeping." Most small business owners just want to know: Am I making money? Where is it going? Am I going to run out? Because software is too complex, they either ignore their finances until tax season or pay expensive consultants.')
    
    p = doc.add_paragraph()
    p.add_run('The Solution: ').bold = True
    p.add_run('CashPilot is a zero-setup, AI-powered financial co-pilot. It acts as a translator between raw bank data and the business owner. Users simply upload a bank statement, and CashPilot\'s AI instantly categorizes transactions, scores their financial health, and explains their business status in plain English. No accounting degree required.')
    
    # Section 2
    doc.add_heading('2. Current State (Hackathon Deliverable)', level=1)
    doc.add_paragraph('Built in 48 hours, the current CashPilot platform is a fully functional, deployed web application featuring:')
    
    doc.add_paragraph('AI Data Pipeline: Automated CSV parsing with GPT-4.1-mini driven transaction categorization (with rule-based fallbacks).', style='List Bullet')
    
    p = doc.add_paragraph('Intelligence Layer:', style='List Bullet')
    doc.add_paragraph('Financial Health Score (0-100): A composite metric evaluating cashflow, expense trends, and revenue diversity.', style='List Bullet 2')
    doc.add_paragraph('Cash Runway Predictor: Calculates how many months the business can survive at its current burn rate.', style='List Bullet 2')
    doc.add_paragraph('Anomaly Detection: Automatically flags unusual spending spikes (e.g., "You spent 3x your average on software this month").', style='List Bullet 2')
    
    doc.add_paragraph('Conversational AI: A chat interface allowing users to ask plain-English questions about their data (e.g., "What were my biggest expenses last month?").', style='List Bullet')
    doc.add_paragraph('Proactive Reporting: Automated, plain-English email summaries sent directly to the user, ensuring they stay informed without needing to log in.', style='List Bullet')
    
    # Section 3
    doc.add_heading('3. Product Roadmap (Next Versions)', level=1)
    doc.add_paragraph('CashPilot will evolve from a reactive analysis tool into a proactive, autonomous financial manager.')
    
    doc.add_heading('v1.5 (Next 30 Days) — Seamless Sync', level=2)
    doc.add_paragraph('Live Bank Feeds: Integration with Plaid or GoCardless to automatically pull daily transactions, eliminating the need for manual CSV uploads.', style='List Bullet')
    doc.add_paragraph('Receipt Capture: Mobile-friendly OCR to snap photos of receipts and auto-match them to bank transactions.', style='List Bullet')
    
    doc.add_heading('v2.0 (Next 90 Days) — Predictive & Actionable', level=2)
    doc.add_paragraph('Tax Prep Mode: AI automatically flags tax-deductible expenses and generates a ready-to-file Schedule C summary for CPAs.', style='List Bullet')
    doc.add_paragraph('Cashflow Forecasting: Predictive modeling that anticipates next month\'s balance based on historical recurring expenses and seasonal income trends.', style='List Bullet')
    doc.add_paragraph('Smart Alerts: SMS/Push notifications for low balance warnings or upcoming large recurring payments.', style='List Bullet')
    
    doc.add_heading('v3.0 (Long Term) — The Autonomous CFO', level=2)
    doc.add_paragraph('Automated Invoicing: Generate and track invoices directly from the chat interface ("Send an invoice to ABC Corp for $500").', style='List Bullet')
    doc.add_paragraph('Expense Optimization: AI identifies unused SaaS subscriptions or cheaper vendor alternatives and offers to cancel/switch them on the user\'s behalf.', style='List Bullet')
    
    # Section 4
    doc.add_heading('4. Business Plan & Monetization', level=1)
    
    doc.add_heading('Target Market', level=2)
    doc.add_paragraph('Our primary market is the 33 million small businesses in the US, specifically targeting the 80% that are "non-employer firms" (solo founders, freelancers, independent contractors). These users are currently underserved by enterprise tools and overcharged by human bookkeepers.')
    
    doc.add_heading('Revenue Model (Freemium SaaS)', level=2)
    doc.add_paragraph('Basic Tier (Free): Manual CSV uploads, basic AI categorization, standard dashboard, 5 AI chat queries per month. Goal: User acquisition and product-led growth.', style='List Bullet')
    doc.add_paragraph('Pro Tier ($12/month or $120/year): Live bank sync (Plaid), unlimited AI chat, anomaly alerts, tax-prep exports, and automated email reporting.', style='List Bullet')
    doc.add_paragraph('Why it works: At $12/month, CashPilot is a fraction of the cost of QuickBooks ($30+/mo) or a human bookkeeper ($200+/mo). It is priced as an impulse buy for a freelancer looking to save 5 hours a month on spreadsheets.', style='List Bullet')
    
    doc.add_heading('Go-to-Market Strategy', level=2)
    doc.add_paragraph('Content Marketing: "How-to" guides for freelancer taxes and cashflow management.', style='List Number')
    doc.add_paragraph('Partnerships: Partnering with freelance platforms (Upwork, Fiverr) and creator economy tools to offer CashPilot as an add-on perk.', style='List Number')
    doc.add_paragraph('Viral Loop: "Share your financial health score" or referral links that grant free months of Pro.', style='List Number')
    
    # Section 5
    doc.add_heading('5. Competitive Advantage (Why We Win)', level=1)
    doc.add_paragraph('Simplicity over Features: We are actively not building double-entry accounting. We are building financial translation. We win by being the easiest tool to use, not the most complex.', style='List Bullet')
    doc.add_paragraph('Conversational Interface: Competitors force users to click through complex dashboards to find answers. CashPilot allows users to simply ask questions in natural language.', style='List Bullet')
    doc.add_paragraph('Proactive, not Reactive: Instead of waiting for users to log in to view a chart, CashPilot pushes plain-English insights and warnings directly to their inbox. We tell them what they need to know before they know to ask.', style='List Bullet')
    
    # Save the document
    doc.save('CashPilot_Executive_Summary.docx')
    print("Successfully created CashPilot_Executive_Summary.docx")

if __name__ == "__main__":
    create_executive_summary()